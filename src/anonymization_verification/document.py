"""This module provides a class for representing a Word document and checking if
it contains gendered language or named entities."""

import json
import logging
import pathlib
from typing import Iterable

import docx
import spacy

from anonymization_verification import config, conjugations

settings = config.get_settings()
SPACY_MODEL = settings.spacy_model
LOGGER_NAME = settings.logger_name
ALLOWED_ENTITIES_FILE = settings.allowed_entities_file
DISALLOWED_WORDS_FILE = settings.disallowed_words_file

logger = logging.getLogger(LOGGER_NAME)
logger.setLevel(logging.INFO)
formatter = logging.Formatter("%(asctime)s - %(name)s - %(levelname)s - %(message)s")
ch = logging.StreamHandler()
ch.setFormatter(formatter)
logger.addHandler(ch)


class AnonyimityVerifier:
    """Class for verifying anonymity of (an iterable of) string(s).

    Attributes:
        text: An iterable of strings or a string.
    """

    def __init__(
        self,
        text: Iterable[str],
        entity_recognition_model: str | None = SPACY_MODEL,
        disallowed_words: Iterable[str] | None = None,
        allowed_entities: Iterable[str] | None = None,
    ) -> None:
        """Initialize the AnoynimityVerifier object.

        Args:
            text: An iterable of strings or a string.
            entity_recognition_model: The name of the spaCy model to use for named entity recognition.
                If None, no model is loaded.
            disallowed_words: An iterable of disallowed words. If None, the default list is used.
            allowed_entities: An iterable of allowed named entities. If None, the default list is used.

        """
        self.text = " ".join(text)
        self.entity_recognition_model = (
            spacy.load(entity_recognition_model) if entity_recognition_model else None
        )

        self.allowed_entities = allowed_entities or self._default_allowed_entities()
        self.disallowed_words = disallowed_words or self._default_disallowed_words()

    def find_any_disallowed(self) -> set[str]:
        """Check if the document passes all checks.

        Returns:
            The words that are not allowed.
        """
        return (
            self.find_disallowed_words()
            | self.find_named_entities()
            | self.find_faulty_conjugations()
        )

    def find_disallowed_words(self) -> set[str]:
        """Check if the document contains disallowed words.

        Returns:
            True if the document contains disallowed words, False otherwise.
        """
        disallowed = {
            word for word in self.text.split() if word.lower() in self.disallowed_words
        }

        return disallowed

    def find_faulty_conjugations(self) -> set[str]:
        """Check if the document contains verbs that are not conjugated properly
        for both he/she and they.

        Returns:
            The verbs that are not conjugated properly.

        """
        english_conjugator = conjugations.CorrectConjugations()
        return english_conjugator.find_faulty_conjugation(self.text)

    def find_named_entities(self) -> set[str]:
        """Find name entities in the document.

        Returns:
            The detected named entities.
        """
        if not self.entity_recognition_model:
            raise ValueError(
                "No spaCy model loaded. Please provide a spaCy model name to the constructor."
            )

        doc = self.entity_recognition_model(self.text)
        entities = {entity.text for entity in doc.ents}
        disallowed = entities - set(self.allowed_entities)

        return disallowed

    @staticmethod
    def _default_allowed_entities() -> set[str]:
        with open(ALLOWED_ENTITIES_FILE, "r", encoding="utf-8") as file_buffer:
            return set(json.load(file_buffer))

    @staticmethod
    def _default_disallowed_words() -> set[str]:
        with open(DISALLOWED_WORDS_FILE, "r", encoding="utf-8") as file_buffer:
            return set(json.load(file_buffer))


class WordDocument(AnonyimityVerifier):
    """A class representing a Word document that can be used to verify the anonymity of its contents.
    Inherits from `AnonyimityVerifier`.

    Args:
        filename: The path to the Word document file.

    Attributes:
        filename: The path to the Word document file.
        document: The Word document object.
    """

    def __init__(
        self,
        filename: str | pathlib.Path,
        entity_recognition_model: str | None = SPACY_MODEL,
    ) -> None:
        """Initialize the WordDocument object.

        Args:
            filename: The path to the Word document file.

        """
        self.filename = pathlib.Path(filename)
        self.document = docx.Document(str(filename))
        paragraphs = []
        for paragraph in self.document.paragraphs:
            if isinstance(paragraph, str):
                paragraphs.append(paragraph)
            else:
                paragraphs.append(paragraph.text)
        super().__init__(paragraphs, entity_recognition_model=entity_recognition_model)

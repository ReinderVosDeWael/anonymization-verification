""" Tests for the document module. """
# pylint: disable=redefined-outer-name
import pathlib
import tempfile
from typing import Generator

import docx
import pytest

from anonymization_verification import document


@pytest.fixture
def word_document(word_text: list[str]) -> Generator[str, None, None]:
    """Creates a Word document with the given text.

    Args:
        word_text: A list of strings representing the paragraphs to be added to
            the document.

    Returns:
        str: The path to the created Word document.
    """
    with tempfile.NamedTemporaryFile(suffix=".docx") as file_buffer:
        doc = docx.Document()
        for paragraph in word_text:
            doc.add_paragraph(paragraph)
        doc.save(file_buffer.name)
        yield str(pathlib.Path(file_buffer.name))


@pytest.mark.parametrize(
    "word_text, expected",
    [
        [["This is a test."], set()],
        [["He/she/they is allowed."], set()],
        [["He is a pronoun."], {"He"}],
        [["Mr. is not allowed"], {"Mr."}],
        [["Ms. is not allowed"], {"Ms."}],
        [["First sentence allowed", "Himself second is not"], {"Himself"}],
    ],
)
def test_check_disallowed_words(word_document: str, expected: set[str]) -> None:
    """Test the check_disallowed_words method."""
    doc = document.WordDocument(word_document, None)

    assert doc.find_disallowed_words() == expected


@pytest.mark.parametrize(
    "word_text, expected",
    [
        [["This is a test."], set()],
        [["He/she/they is allowed."], set()],
        [["Mcdonalds is not allowed"], {"Mcdonalds"}],
        [["This sentence is allowed", "Sir, this is a Subway"], {"Subway"}],
        [["Child Mind Institute is allowed"], set()],
    ],
)
def test_check_named_entities(word_document: str, expected: set[str]) -> None:
    """Test the check_named_entities method."""
    doc = document.WordDocument(word_document, "en_core_web_sm")

    assert doc.find_named_entities() == expected
